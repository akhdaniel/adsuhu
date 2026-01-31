#!/usr/bin/python
#-*- coding: utf-8 -*-

from odoo import models, fields, api, _
from odoo.exceptions import UserError
import html
import json
import logging
_logger = logging.getLogger(__name__)
import re
import requests
import time
import base64
import io
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT, WD_SECTION
from odoo.modules.module import get_module_path
from bs4 import BeautifulSoup
# from docx2pdf import convert
from .libs.openai_lib import *

try:
    from bs4 import BeautifulSoup
except Exception:  # pragma: no cover - optional dependency
    BeautifulSoup = None

try:
    import html2text as html2text_lib
except Exception:  # pragma: no cover - optional dependency
    html2text_lib = None

_logger = logging.getLogger(__name__)


DEFAULT_SPECIFIC_INSTRUCTION="""REQUIRED JSON OUTPUT FORMAT:

{
  "product": "...",
  "category": "...",
  "level_maslow": "...",
  "unique_selling_propositions": [
    "...",
    "...
    "...",
    "...",
    "..."
  ],
  "value_map_extended": [
    {
      "feature": "...",
      "pain_point": "...",
      "gain_point": "...",
      "functional_benefits": "...",
      "emotional_benefits": "...",
      "proof": "...",
      "buying_motif": "...",
      "buying_triggers": "...",
      "level_maslow": "...",
      "relevan_usp": "..."
    },
    ... # list semua feature 
  ],
  "differentiation_spike": "...",
  "buying_triggers": {
    "rational": [
      "...",
      "...",
      "..."
    ],
    "emotional": [
      "...",
      "...",
      "..."
    ]
  },
  "initial_target_market": {
    "persona": "...",
    "pain": "...",
    "gain": "..."
  },
  "copywriting_angle": {
    "hook": "...",
    "proof": "....",
    "path": "..."
  }
}

"""

class product_value_analysis(models.Model):
    _name = "vit.product_value_analysis"
    _inherit = "vit.product_value_analysis"
    specific_instruction = fields.Text( string=_("Specific Instruction"), default=DEFAULT_SPECIFIC_INSTRUCTION)

    def _get_default_lang(self):
        return self.env["res.lang"].search(
            [("iso_code", "=", "en")], limit=1
        ).id
    lang_id = fields.Many2one(comodel_name="res.lang", default=_get_default_lang)

    def _get_default_gpt_model(self):
        return self.env["vit.gpt_model"].search(
            [("name", "=", "deepseek-chat")], limit=1
        ).id
    gpt_model_id = fields.Many2one(comodel_name="vit.gpt_model", default=_get_default_gpt_model)

    def _get_default_prompt(self):
        prompt = self.env.ref("vit_ads_suhu_inherit.gpt_product_value_analysis", raise_if_not_found=False)
        if prompt:
            return prompt.id
        return self.env["vit.gpt_prompt"].search(
            [("name", "=", "product_value_analysis")], limit=1
        ).id


    def _get_default_write_prompt(self):
        prompt = self.env.ref("vit_ads_suhu_inherit.write_description_prompt", raise_if_not_found=False)
        if prompt:
            return prompt.id
        return self.env["vit.gpt_prompt"].search(
            [("name", "=", "write_description_prompt")], limit=1
        ).id

    gpt_prompt_id = fields.Many2one(comodel_name="vit.gpt_prompt",  string=_("Analyzer Prompt"), default=_get_default_prompt)
    write_gpt_prompt_id = fields.Many2one(comodel_name="vit.gpt_prompt",  string=_("Description Writer Prompt"), default=_get_default_write_prompt)

    # write description and features
    def action_write_with_ai(self, ):
        # def json_to_markdown(features):
        #     output = []
        #     for feature in features:
        #         for title, items in feature.items():
        #             output.append(f"## {title}")
        #             for item in items:
        #                 output.append(f"- {item}")
        #             output.append("")  # blank line between features
        #     return "\n".join(output)     
           
        if not self.write_gpt_prompt_id:
            raise UserError('Write GPT prompt empty')
        if not self.gpt_model_id:
            raise UserError('Write GPT model empty')
        
        context = self.description
        additional_command=f"Reponse in {self.lang_id.name}"
        system_prompt = self.write_gpt_prompt_id.system_prompt 
        question = ""
        user_prompt = self.write_gpt_prompt_id.user_prompt

        openai_api_key = self.env["ir.config_parameter"].sudo().get_param("deepseek_api_key")
        openai_base_url = self.env["ir.config_parameter"].sudo().get_param("deepseek_base_url", None)

        model = self.gpt_model_id.name

        response = generate_content(openai_api_key=openai_api_key, 
                                openai_base_url=openai_base_url, 
                                model=model, 
                                system_prompt=system_prompt, 
                                user_prompt=user_prompt, 
                                context=context, 
                                question=question, 
                                additional_command=additional_command)    
        response = self.clean_md(response)
        _logger.info(response)

        response = json.loads(response)
        if not 'description' in response:
            raise UserError(f'Error from GPT: {response}')
        self.description = response.get('description','')
        # self.features = json_to_markdown(response['features'])
        self.features = self.md_to_html( 
            self.json_to_markdown(
                response.get('features') 
            )            
        )


    @api.onchange("description","features","lang_id","specific_instruction")
    def _get_input(self, ):
        """
        {
        }
        """
        for rec in self:
            rec.input = f"""

# DESCRIPTION:
---
{rec.description}

# PRODUCT FEATURES:
---
{rec.features}

# INSTRUCTIONS:
---
{self.general_instruction}

{self.specific_instruction or ''}

Response in {self.lang_id.name} language.

"""


    def action_fetch_product_description(self, ):
        for record in self:
            if not record.product_url:
                raise UserError(_("Please provide a product URL before fetching the description."))

            html_content = record._get_product_content(record.product_url)
            cleaned_html = record._clean_html_content(html_content)
            markdown_content = record._html_to_markdown(cleaned_html)
            record.description = markdown_content or cleaned_html or ""

        return True


    def action_generate(self, ):

        if not self.gpt_prompt_id:
            raise UserError('Analyze GPT prompt empty')
        if not self.gpt_model_id:
            raise UserError('GPT model empty')

        openai_api_key = self.env["ir.config_parameter"].sudo().get_param("deepseek_api_key")
        openai_base_url = self.env["ir.config_parameter"].sudo().get_param("deepseek_base_url", None)

        model = self.gpt_model_id.name

        user_prompt = self.gpt_prompt_id.user_prompt
        user_prompt += f"{self.input}\n"
        system_prompt = self.gpt_prompt_id.system_prompt 

        context = ""
        additional_command=""
        question = ""

        response = generate_content(openai_api_key=openai_api_key, 
                                openai_base_url=openai_base_url, 
                                model=model, 
                                system_prompt=system_prompt, 
                                user_prompt=user_prompt, 
                                context=context, 
                                question=question, 
                                additional_command=additional_command)    

        response = self.clean_md(response)
        self.output = response

        self.generate_output_html()
        

    def _get_product_content(self, url):
        if "shopee." in url:
            shopee_content = self._fetch_shopee_product_content(url)
            if shopee_content:
                return shopee_content
        return self._fetch_product_page(url)

    def _fetch_product_page(self, url):
        try:
            response = requests.get(url, timeout=20)
            response.raise_for_status()
            return response.text
        except requests.RequestException as exc:
            _logger.exception("Failed to download product page %s", url)
            raise UserError(
                _("Unable to fetch the product page.\n\n%(error)s") % {"error": str(exc)}
            )

    def _clean_html_content(self, html_text):
        if not html_text or not BeautifulSoup:
            return html_text
        soup = BeautifulSoup(html_text, "html.parser")
        for tag in soup(["script", "style", "noscript"]):
            tag.decompose()
        body = soup.body or soup
        return body.prettify()

    def _html_to_markdown(self, html_text):
        if not html_text:
            return ""
        if html2text_lib:
            converter = html2text_lib.HTML2Text()
            converter.ignore_links = False
            converter.body_width = 0
            return converter.handle(html_text)
        if BeautifulSoup:
            soup = BeautifulSoup(html_text, "html.parser")
            return soup.get_text("\n", strip=True)
        return html_text

    # Shopee scraper
    def _fetch_shopee_product_content(self, url):
        ids = self._extract_shopee_ids(url)
        if not ids:
            _logger.info("Unable to extract Shopee identifiers from url %s", url)
            return ""
        shop_id, item_id = ids
        product_data = self._call_shopee_api(shop_id, item_id, url)
        if product_data:
            return self._build_shopee_product_html(product_data)

        html_page = self._fetch_product_page(url)
        structured_html = self._extract_shopee_content_from_html(html_page)
        return structured_html or html_page or ""

    def _call_shopee_api(self, shop_id, item_id, referer_url):
        api_url = f"https://shopee.co.id/api/v4/item/get?itemid={item_id}&shopid={shop_id}"
        headers = {
            "User-Agent": "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 "
            "(KHTML, like Gecko) Chrome/123.0.0.0 Safari/537.36",
            "Accept": "application/json, text/javascript, */*; q=0.01",
            "Accept-Language": "en-US,en;q=0.9",
            "Referer": referer_url,
            "X-Requested-With": "XMLHttpRequest",
            "X-API-SOURCE": "pc",
        }
        try:
            response = requests.get(api_url, headers=headers, timeout=20)
            if response.status_code == 403:
                _logger.warning("Shopee API request blocked with 403 for %s", api_url)
                return None
            response.raise_for_status()
            payload = response.json()
        except requests.RequestException as exc:
            _logger.warning("Shopee API request failed: %s", exc)
            return None
        except ValueError:
            _logger.warning("Shopee API returned non JSON response for %s", api_url)
            return None

        item = payload.get("data") or payload.get("item") or {}
        if not item:
            _logger.info("Shopee API returned empty item for url %s", url)
            return None
        return item

    def _build_shopee_product_html(self, item):
        title = item.get("name") or ""
        description = (item.get("description") or "").replace("\r\n", "\n")
        attributes = item.get("attributes") or []
        price = item.get("price") or item.get("price_min") or item.get("price_max") or ""

        html_parts = []
        if title:
            html_parts.append(f"<h1>{html.escape(title)}</h1>")
        if price:
            if isinstance(price, (int, float)):
                price = price / 100000
                price_display = f"Rp {price:,.0f}"
            else:
                price_display = str(price)
            html_parts.append(f"<p><strong>Price:</strong> {html.escape(price_display)}</p>")
        if description:
            paragraphs = "".join(
                f"<p>{html.escape(line)}</p>" for line in description.split("\n") if line.strip()
            )
            html_parts.append(paragraphs)
        if attributes:
            html_parts.append("<ul>")
            for attr in attributes:
                name = attr.get("name") or attr.get("attribute_name") or ""
                value_list = attr.get("values") or attr.get("attribute_value") or []
                if isinstance(value_list, list):
                    parsed_values = []
                    for val in value_list:
                        if isinstance(val, dict):
                            parsed_values.append(val.get("name") or val.get("value") or "")
                        else:
                            parsed_values.append(str(val))
                    value = ", ".join(filter(None, parsed_values))
                else:
                    value = value_list or ""
                if name or value:
                    html_parts.append(
                        f"<li><strong>{html.escape(name)}:</strong> {html.escape(str(value))}</li>"
                    )
            html_parts.append("</ul>")
        return "".join(html_parts)

    def _extract_shopee_content_from_html(self, html_text):
        if not html_text:
            return ""
        if not BeautifulSoup:
            return html_text
        soup = BeautifulSoup(html_text, "html.parser")
        scripts = soup.find_all("script", {"type": "application/ld+json"})
        for script in scripts:
            raw = script.string
            if not raw:
                continue
            try:
                data = json.loads(raw)
            except ValueError:
                continue
            if isinstance(data, list):
                product_data = next((item for item in data if isinstance(item, dict) and item.get("@type") == "Product"), None)
            elif isinstance(data, dict) and data.get("@type") == "Product":
                product_data = data
            else:
                product_data = None
            if not product_data:
                continue
            title = product_data.get("name") or ""
            description = product_data.get("description") or ""
            offers = product_data.get("offers") or {}
            price = offers.get("price")
            attributes = product_data.get("additionalProperty") or []

            html_parts = []
            if title:
                html_parts.append(f"<h1>{html.escape(title)}</h1>")
            if price:
                html_parts.append(f"<p><strong>Price:</strong> {html.escape(str(price))}</p>")
            if description:
                html_parts.append(f"<div>{description}</div>")
            if attributes:
                html_parts.append("<ul>")
                for prop in attributes:
                    name = prop.get("name")
                    value = prop.get("value")
                    if name or value:
                        html_parts.append(
                            f"<li><strong>{html.escape(str(name or ''))}:</strong> {html.escape(str(value or ''))}</li>"
                        )
                html_parts.append("</ul>")
            if html_parts:
                return "".join(html_parts)
        return ""

    def _extract_shopee_ids(self, url):
        patterns = [
            r"/product/(\d+)/(\d+)",
            r"-i\.(\d+)\.(\d+)",
        ]
        for pattern in patterns:
            match = re.search(pattern, url)
            if match:
                return match.group(1), match.group(2)
        return None


    def action_generate_report(self, ):

        
        product = self
        report = []

        # ------------------------------------------------------------------------
        # 1. Product value
        # ------------------------------------------------------------------------
        report.append("# PRODUCT DESCRIPTION")
        report.append("## Description")
        report.append(f"{product.description}")
        report.append("")
        report.append(f"Product URL: {product.product_url or '-'} ")
        report.append(f"Client Name: {product.partner_id.name}")
        report.append("")        
        report.append("## Features")
        report.append("---")
        report.append(f"{product.features}")
        report.append("")
        report.append("--- SECTION 1 PRODUCT VALUE ---")
        report.append("# PRODUCT VALUE ANALYSIS")
        report.append("---")
        output = json.loads(self.clean_md(product.output))

        report.append(f"## Product category")
        report.append(f"{output['category']}")
        report.append("\n")

        report.append(f"## Level Maslow")
        report.append(f"{output['level_maslow']}")
        report.append("\n")

        report.append(f"## Unique Selling Propositions")
        report.append(self.list_to_bullet(output['unique_selling_propositions']))
        
        report.append(f"## Extended Value Map")
        for i,val in enumerate(output['value_map_extended']):
            report.append(f"### Feature {i+1}: {val['feature']}")
            report.append("\n")
            report.append(f"* Pain Point: {val['pain_point']}")
            report.append(f"* Gain Point: {val['gain_point']}")
            report.append(f"* Functional Benefits: {val['functional_benefits']}")
            report.append(f"* Emotional Benefits: {val['emotional_benefits']}")
            report.append(f"* Proof: {val['proof']}")
            report.append(f"* Buying Motif: {val['buying_motif']}")
            report.append(f"* Buying Trigger: {val['buying_triggers']}")
            report.append(f"* Level Maslow: {val['level_maslow']}")
            report.append(f"* Relevan USPs: {val['relevan_usp']}")
        report.append("\n")
        
        report.append("## Differentiation Spike")
        report.append(output['differentiation_spike'])
        report.append("\n")
        
        report.append("## Buying Triggers")
        buying_triggers=output['buying_triggers']
        report.append("### Rational")
        report.append(self.list_to_bullet(buying_triggers['rational']))
        report.append("\n")
        
        report.append("### Emotional")
        report.append(self.list_to_bullet(buying_triggers['emotional']))
        report.append("\n")       

        report.append("## Initial Target Market")
        initial_target_market = output['initial_target_market']
        report.append(f"* *Persona*: {initial_target_market['persona']}")
        report.append(f"* *Pain*: {initial_target_market['pain']}")
        report.append(f"* *Gain*: {initial_target_market['gain']}")
        report.append("\n")

        report.append("## Copywriting Angles")
        copywriting_angle = output['copywriting_angle']
        report.append(f"* *Hook*: {copywriting_angle['hook']}")
        report.append(f"* *Proof*: {copywriting_angle['proof']}")
        report.append(f"* *Path*: {copywriting_angle['path']}")
        report.append("\n")

        # ------------------------------------------------------------------------
        # 2. Market analysis
        # ------------------------------------------------------------------------
        report.append("--- SECTION 2 MARKET ANALYSIS AND AUDIENCE PROFILE ---")        
        markets = product.market_mapper_ids
        for m, market in enumerate(markets, start=2):
            report.append(f"# MARKET ANALYSIS")
            report.append("---")
            res = self.json_to_markdown(json.loads(self.clean_md(market.output)), level=2, max_level=3, prefix=m)
            report.append(res)
            report.append("\n")
            
            # ------------------------------------------------------------------------
            # 3..p Audience profile
            # ------------------------------------------------------------------------
            profiles = market.audience_profiler_ids
            for p, profile in enumerate(profiles, start=m+1):
                report.append(f"# AUDIENCE PROFILE {p-2}: {profile['description']}")
                report.append("---")

                if not profile.output:
                    report.append("--no data--")
                    continue

                res = self.json_to_markdown(json.loads(self.clean_md(profile.output)), level=2, max_level=3, prefix=p)
                report.append(res)
                report.append("\n")
                report.append("\n")

            # ------------------------------------------------------------------------
            # p+a ... p+(an) Angles
            # ------------------------------------------------------------------------
            report.append("\n")
            report.append("--- SECTION 3 ANGLES ---")        
            profiles = market.audience_profiler_ids
            for p, profile in enumerate(profiles):
                angles = profile.angle_hook_ids
                if not angles:
                    report.append(f"# NO ANGLES - {profile['name']}")
                    report.append("--no angles--")
                    continue
                
                for a, angle in enumerate(angles, start=1):
                    report.append(f"# {angle['name']} - {profile['name']}")
                    report.append("---")                        

                    if not angle.description:
                        report.append("--no description--")
                        continue
                    
                    if not angle.output:
                        report.append("--no data--")
                        continue

                    report.append("## Audience Profile")
                    report.append(profile['name'])
                    report.append(profile['description'])

                    js = json.loads(self.clean_md(angle.output))
                    if a > 1:
                        js.pop('big_ideas')
                        js.pop('strategic_notes')
                    
                    res = self.json_to_markdown(js, level=2, max_level=4, prefix=a)
                    # show big ideas only on first angle
                    
                    report.append(res)
                    report.append("\n")    
                    report.append("\n")    

                    # ------------------------------------------------------------------------
                    # Hooks
                    # ------------------------------------------------------------------------
                    # report.append("## Hooks")
                    # hooks = angle.hook_ids
                    # if not hooks:
                    #     report.append("--no hooks--")
                    #     continue
                    # for h, hook in enumerate(hooks, start=1):
                    #     report.append(f"### Hook {hook['hook_no']}: {hook['description']}")
                    #     report.append("---")

                    #     if not hook.output:
                    #         report.append("--no hook data--")
                    #         continue

                    #     js = json.loads(self.clean_md(hook.output))
                    #     res = json_to_markdown(js['hook'],level=4, max_level=4)
                    #     report.append(res)
                    #     report.append("\n")                


            # ------------------------------------------------------------------------
            # Ads copy per market, audience profile, angle, hooks
            # ------------------------------------------------------------------------
            report.append("\n")
            report.append("--- SECTION 4 ADS COPY ---")        
            ads_count = 1
            profiles = market.audience_profiler_ids
            for p, profile in enumerate(profiles):
                angles = profile.angle_hook_ids
                for a, angle in enumerate(angles, start=1):
                    hooks = angle.hook_ids
                    for h, hook in enumerate(hooks, start=1):
                        ads = hook.ads_copy_ids
                        if not ads:
                            report.append("--no ads--")
                            continue

                        for adx, ad in enumerate(ads, start=1):
                            report.append(f"# ADS {ads_count}: {ad['hook']}")
                            report.append("---")

                            if not ad.output:
                                report.append("--no ads data--")
                                continue

                            report.append(f"## Audience Profile: {profile['audience_profile_no']}")
                            report.append(f"{profile['name']} - {profile['description']}")
                            report.append(f"## Angle: {angle['angle_no']}")
                            report.append(f"{angle['name']} - {angle['description']}")
                            report.append(f"## Hook: {hook['hook_no']} ")
                            report.append(f"{hook['name']} - {hook['description']}")
                                                     
                            js = json.loads(self.clean_md(ad.output))
                            js.pop('angle')
                            js.pop('hook')
                            js.pop('ads_copy')
                            js.pop('landing_page')
                            res = self.json_to_markdown(js, level=2, max_level=3, prefix=a)
                            report.append(res)
                            report.append("\n")  
                            ads_count += 1

                            # ------------------------------------------------------------------------
                            # Ads images
                            # ------------------------------------------------------------------------  
                            images = ad.image_generator_ids
                            if not images:
                                report.append("--no images --")
                                continue

                            for imgx, img in enumerate(images, start=1):
                                js = json.loads(self.clean_md(img.output))
                                js.pop('angle_library')
                                js.pop('hook_library')
                                js.pop('instruction')
                                report.append(f"## {img['name']}")
                                res = self.json_to_markdown(js, level=3, max_level=4).replace('\n\n','\n')
                                report.append(res)
                                report.append("\n")  

                                variants=img.image_variant_ids
                                if not variants:
                                    report.append("--no image variants--")  
                                    continue
                                for imgvx, variant in enumerate(variants, start=1):
                                    report.append(f"### Image Variant: {variant['name']}")
                                    src = f"/web/image/vit.image_variant/{variant.id}/image_1024?unique={int(time.time())}"
                                    res = f"![{ad['name'] or 'image'}]({src})"
                                    report.append(res)
                                    report.append("\n")    

            # ------------------------------------------------------------------------
            # Landing pages per ads copy...
            # ------------------------------------------------------------------------   
            report.append("\n")                         
            report.append("--- SECTION 5 LANDING PAGES ---")        
            lps_count = 1
            profiles = market.audience_profiler_ids
            for p, profile in enumerate(profiles):
                angles = profile.angle_hook_ids
                for a, angle in enumerate(angles, start=1):
                    hooks = angle.hook_ids
                    for h, hook in enumerate(hooks, start=1):
                        ads = hook.ads_copy_ids
                        if not ads:
                            report.append("--no LP--")
                            continue

                        for adx, ad in enumerate(ads, start=1):

                            if not ad.output:
                                report.append("--no ads data--")
                                continue
                            # ------------------------------------------------------------------------
                            # Landing pages
                            # ------------------------------------------------------------------------  
                            lps = ad.landing_page_builder_ids
                            if not lps:
                                report.append("--no LP--")
                                continue
                            
                            for lp in lps:
                                report.append(f"# LP {lps_count}: {ad['hook']}")
                                report.append("---")                                
                                js = json.loads(self.clean_md(lp.output))
                                res = self.json_to_markdown(js, level=2, max_level=3)
                                report.append(res)
                                report.append("\n")  
                                lps_count+=1

        self.final_report = "\n".join(report)


    def add_html_to_docx(self, doc, html_content):
        is_inside_tag = False
        def target_image_width():
            """Compute 75% of usable page width for images."""
            try:
                section = doc.sections[0]
                usable_width = section.page_width - section.left_margin - section.right_margin
                return int(usable_width * 0.75)
            except Exception:
                return None

        def add_inline_runs(paragraph, element, default_bold=False):
            """Render inline children of an element into a paragraph."""
            for child in element.children:
                run = None
                if isinstance(child, str):
                    if not child:
                        continue
                    run = paragraph.add_run(child)
                    if default_bold:
                        run.bold = True
                elif child.name == 'strong':
                    run = paragraph.add_run(child.get_text())
                    run.bold = True
                elif child.name == 'code':
                    run = paragraph.add_run(child.get_text())
                    run.bold = True
                    run.underline = True
                elif child.name == 'em':
                    run = paragraph.add_run(child.get_text())
                    run.italic = True
                elif child.name == 'u':
                    run = paragraph.add_run(child.get_text())
                    run.underline = True
                elif child.name == 'a':
                    href = child.get("href", "")
                    run = paragraph.add_run(f'{child.get_text()} ({href})' if href else child.get_text())
                    run.underline = True
                elif child.name == 'br':
                    paragraph.add_run("\n")
                elif child.name == 'img':
                    img_bytes, resolved_src = fetch_image_bytes(child.get('src'))
                    if img_bytes:
                        try:
                            run = paragraph.add_run()
                            width = target_image_width()
                            if width:
                                run.add_picture(io.BytesIO(img_bytes), width=width)
                            else:
                                run.add_picture(io.BytesIO(img_bytes))
                            link_run = paragraph.add_run(f"\n{resolved_src}")
                            link_run.underline = True
                        except Exception:
                            run = paragraph.add_run(f"[Image could not be embedded: {child.get('alt', '')}]")
                    else:
                        run = paragraph.add_run(f"[Image missing: {child.get('alt', '') or child.get('src', '')}]")
                else:
                    run = paragraph.add_run(child.get_text())
                if default_bold and run:
                    run.bold = True

        def process_li_element(doc, li, is_ordered_list):
            """Process each <li> element, handling nested <strong> and other elements."""
            if is_ordered_list:
                p = doc.add_paragraph(style='List Number')
            else:
                p = doc.add_paragraph(style='List Bullet')

            for sub_element in li.children:
                if sub_element.name == 'strong':
                    run = p.add_run(sub_element.get_text())
                    run.bold = True
                elif sub_element.name == 'code':
                    run = p.add_run(sub_element.get_text())
                    run.underline = True
                    run.bold = True
                elif sub_element.name in ['ol','ul']:
                    for index, li2 in enumerate(sub_element.find_all('li'), start=1):
                        process_li_element(doc, li2, is_ordered_list=False)
                else:
                    run = p.add_run(sub_element.get_text())


        def process_p_element(doc, p_element):
            """Process a <p> element, handling inline tags like <strong>, <em>, etc."""
            p = doc.add_paragraph()  # Create a new paragraph for the <p> element

            add_inline_runs(p, p_element)

        def fetch_image_bytes(src):
            """Fetch image content from a src; supports data URLs and HTTP/HTTPS.

            Returns a tuple of (content_bytes or None, resolved_src_url).
            """
            resolved_src = src
            if not src:
                return None, resolved_src
            if src.startswith('data:image'):
                try:
                    header, encoded = src.split(',', 1)
                    return base64.b64decode(encoded), resolved_src
                except Exception:
                    return None, resolved_src
            if src.startswith('/'):
                try:
                    base_url = self.env['ir.config_parameter'].sudo().get_param('web.base.url') or ''
                except Exception:
                    base_url = ''
                if base_url:
                    resolved_src = base_url.rstrip('/') + src
            if resolved_src.startswith('http://') or resolved_src.startswith('https://'):
                try:
                    resp = requests.get(resolved_src, timeout=10)
                    resp.raise_for_status()
                    return resp.content, resolved_src
                except Exception:
                    return None, resolved_src
            try:
                with open(resolved_src, 'rb') as f:
                    return f.read(), resolved_src
            except Exception:
                return None, resolved_src

        def process_table(doc, table_element):
            """Render HTML table into docx table."""
            rows = table_element.find_all('tr')
            if not rows:
                return

            col_count = max((len(r.find_all(['th', 'td'])) for r in rows), default=0)
            if not col_count:
                return

            table = doc.add_table(rows=len(rows), cols=col_count)
            self.ensure_table_style(doc, 'Grid Table 4 - Accent 1')
            try:
                table.style = 'Grid Table 4 - Accent 1'
            except KeyError:
                if 'Table Grid' in doc.styles:
                    table.style = 'Table Grid'

            for row_idx, row in enumerate(rows):
                cells = row.find_all(['th', 'td'])
                for col_idx in range(col_count):
                    if col_idx >= len(cells):
                        continue
                    cell_element = cells[col_idx]
                    cell = table.rows[row_idx].cells[col_idx]
                    cell.text = ""
                    used_first_paragraph = False

                    def next_paragraph(style_name=None):
                        nonlocal used_first_paragraph
                        if not used_first_paragraph and cell.paragraphs:
                            used_first_paragraph = True
                            para = cell.paragraphs[0]
                            if style_name:
                                para.style = style_name
                            return para
                        used_first_paragraph = True
                        return cell.add_paragraph(style=style_name) if style_name else cell.add_paragraph()

                    bold_default = cell_element.name == 'th'
                    has_child_content = False
                    for child in cell_element.children:
                        if isinstance(child, str):
                            if not child.strip():
                                continue
                            para = next_paragraph()
                            run = para.add_run(child)
                            if bold_default:
                                run.bold = True
                            has_child_content = True
                        elif child.name == 'p':
                            para = next_paragraph()
                            add_inline_runs(para, child, default_bold=bold_default)
                            has_child_content = True
                        elif child.name in ['ul', 'ol']:
                            style = 'List Number' if child.name == 'ol' else 'List Bullet'
                            first_li = True
                            for li in child.find_all('li', recursive=False):
                                if first_li:
                                    para = next_paragraph(style)
                                    first_li = False
                                else:
                                    para = cell.add_paragraph(style=style)
                                add_inline_runs(para, li, default_bold=bold_default)
                            has_child_content = True
                        elif child.name == 'img':
                            img_bytes, resolved_src = fetch_image_bytes(child.get('src'))
                            para = next_paragraph()
                            if img_bytes:
                                try:
                                    run = para.add_run()
                                    width = target_image_width()
                                    if width:
                                        run.add_picture(io.BytesIO(img_bytes), width=width)
                                    else:
                                        run.add_picture(io.BytesIO(img_bytes))
                                    link_run = cell.add_paragraph().add_run(resolved_src)
                                    link_run.underline = True
                                except Exception:
                                    para.add_run(f"[Image could not be embedded: {child.get('alt', '')}]")
                            else:
                                para.add_run(f"[Image missing: {child.get('alt', '') or child.get('src', '')}]")
                            has_child_content = True
                        else:
                            para = next_paragraph()
                            add_inline_runs(para, child, default_bold=bold_default)
                            has_child_content = True

                    if not has_child_content:
                        para = next_paragraph()
                        run = para.add_run(cell_element.get_text())
                        if bold_default:
                            run.bold = True


        # Parse the HTML content
        soup = BeautifulSoup(html_content, 'html.parser')
        ol_count = 0  # Counter to track ordered lists and reset numbering
        for element in soup.descendants:
            if element.name == 'p' and element.parent.name == '[document]':
                section_match = re.match(r'^\s*---\s*SECTION\s+(.*?)\s*---\s*$', element.get_text().strip(), re.IGNORECASE)
                if section_match:
                    section_title = section_match.group(1).strip()
                    doc.add_page_break()
                    doc.add_paragraph(f"SECTION {section_title}", style='Title')
                    doc.add_page_break()
                    continue

            if element.name == 'h1':
                doc.add_heading(element.get_text(), level=1)
            elif element.name == 'h2':
                doc.add_heading(element.get_text(), level=2)
            elif element.name == 'h3':
                doc.add_heading(element.get_text(), level=3)
            elif element.name == 'h4':
                doc.add_heading(element.get_text(), level=4)
            elif element.name == 'h5':
                doc.add_heading(element.get_text(), level=5)
            elif element.name == 'h6':
                doc.add_heading(element.get_text(), level=6)

            elif element.name == 'table':
                if element.parent.name!='[document]':
                    continue
                process_table(doc, element)

            elif element.name == 'img':
                src = element.get('src')
                alt = element.get('alt', '')
                img_bytes, resolved_src = fetch_image_bytes(src)
                if img_bytes:
                    try:
                        width = target_image_width()
                        if width:
                            doc.add_picture(io.BytesIO(img_bytes), width=width)
                        else:
                            doc.add_picture(io.BytesIO(img_bytes))
                        p = doc.add_paragraph()
                        link_run = p.add_run(resolved_src)
                        link_run.underline = True
                    except Exception:
                        p = doc.add_paragraph()
                        p.add_run(f"[Image could not be embedded: {alt or src}]")
                else:
                    p = doc.add_paragraph()
                    p.add_run(f"[Image missing: {alt or src}]")

            elif element.name == 'p' and '```' not in element.get_text():
                if element.parent.name!='[document]':
                    continue
                process_p_element(doc, element)

            elif element.name == 'strong':
                if element.parent.name!='[document]':
                    continue
                p = doc.add_paragraph()
                run = p.add_run(element.get_text())
                run.bold = True

            elif element.name == 'em':
                if element.parent.name!='[document]':
                    continue
                p = doc.add_paragraph()
                run = p.add_run(element.get_text())
                run.italic = True

            elif element.name in ['ul','ol']:
                if element.parent.name!='[document]':
                    continue
                for index, li in enumerate(element.find_all('li'), start=1):
                    process_li_element(doc, li, is_ordered_list=False)

            elif element.name == 'code' or '```' in element.get_text():  # and element.parent.name == 'pre':
                if element.parent.name!='[document]':
                    continue

                codes = element.get_text()
                for code in codes.split('\n'):
                    code = code.strip()
                    if code in ['bash','sh', 'ini','python','sql','menu', 'csv', 'txt', 'cmd']:
                        p = doc.add_paragraph(code, style="Coding Head")
                    else:
                        p = doc.add_paragraph(code, style="Coding")

            elif element.name == 'a':
                if element.parent.name!='[document]':
                    continue

                p = doc.add_paragraph()
                if 'href' in element:
                    run = p.add_run(f'{element.get_text()} ({element["href"]})')
                else:
                    run = p.add_run(f'{element.get_text()}')
                run.underline = True


    def strip_html_tags(self, text):
        soup = BeautifulSoup(text, "html.parser")
        return soup.get_text()



    # Function to convert HTML to DOCX
    def html_to_docx(self, html_content, docx_file_path):
        soup = BeautifulSoup(html_content, 'html.parser')
        doc = Document()

        # Process each HTML tag and add it to DOCX
        for tag in soup.find_all(True):  # Finds all HTML tags
            if tag.name == 'h1':
                doc.add_heading(tag.get_text(), level=1)
            elif tag.name == 'h2':
                doc.add_heading(tag.get_text(), level=2)
            elif tag.name == 'h3':
                doc.add_heading(tag.get_text(), level=3)
            elif tag.name == 'p':
                doc.add_paragraph(tag.get_text())
            elif tag.name == 'li':
                doc.add_paragraph(tag.get_text(), style='ListBullet')

        # Save the DOCX file
        doc.save(docx_file_path)


    def ensure_style_exists(self, doc, style_name, style_type):
        if style_name not in doc.styles:
            new_style = doc.styles.add_style(style_name, style_type)
            if style_type == WD_STYLE_TYPE.PARAGRAPH:
                new_style.base_style = doc.styles['Normal']
        return doc.styles[style_name]

    def ensure_table_style(self, doc, style_name, base_style_name='Table Grid'):
        """Ensure a table style exists so formatting can be applied safely."""
        styles = doc.styles
        if style_name in styles:
            return styles[style_name]
        new_style = styles.add_style(style_name, WD_STYLE_TYPE.TABLE)
        try:
            new_style.base_style = styles[base_style_name]
        except KeyError:
            # If base style is missing, leave it with default formatting
            pass
        return new_style

    def setup_heading_styles(self, doc):
        styles = doc.styles

        # Heading 1 style
        heading1_style = styles.add_style('Custom Heading 1', WD_STYLE_TYPE.PARAGRAPH)
        heading1_style.base_style = styles['Heading 1']
        heading1_style.font.size = Pt(30)
        heading1_style.font.bold = True
        heading1_format = heading1_style.paragraph_format
        heading1_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        heading1_format.space_before = Pt(12)
        heading1_format.space_after = Pt(6)
        heading1_format.keep_with_next = True

        # Heading 2 style
        heading2_style = styles.add_style('Custom Heading 2', WD_STYLE_TYPE.PARAGRAPH)
        heading2_style.base_style = styles['Heading 2']
        heading2_style.font.size = Pt(14)
        heading2_style.font.bold = True
        heading2_format = heading2_style.paragraph_format
        heading2_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        heading2_format.space_before = Pt(10)
        heading2_format.space_after = Pt(6)
        heading2_format.keep_with_next = True

        return heading1_style, heading2_style


    def action_download_docx(self, ):

        self.action_generate_report()
        cover = self.report_template
        doc = None
        if cover:
            decoded_bytes = base64.b64decode(cover)
            cover_file = io.BytesIO(decoded_bytes)
            doc = Document(cover_file)
        else:
            doc = Document()

        style1 = self.ensure_style_exists(doc, 'Subtitle', WD_STYLE_TYPE.PARAGRAPH)
        style2 = self.ensure_style_exists(doc, 'List Bullet', WD_STYLE_TYPE.PARAGRAPH)
        style3 = self.ensure_style_exists(doc, 'List Number', WD_STYLE_TYPE.PARAGRAPH)
        coding = self.ensure_style_exists(doc, 'Coding', WD_STYLE_TYPE.PARAGRAPH)

        # heading1_style, heading2_style = self.setup_heading_styles(doc)
        # Title
        title = self.name
        doc.add_paragraph(f'{title}', style='Title')
        doc.add_paragraph(self.description, style='Subtitle' )

        # Introduction
        # doc.add_heading('Summary', level=1)
        # self.add_html_to_docx(doc, self.summary)

        # Loop content
        html_content = self.md_to_html(self.final_report)
        html_content = html_content.replace('</code></p>', '</code>').replace('<p><code>', '<code>')
        html_content = re.sub(r'<li>\s*<p>(.*?)</p>\s*</li>', r'<li>\1</li>', html_content, flags=re.DOTALL)

        self.add_html_to_docx(doc, html_content)


        # save file
        tmp = self.name.replace('/','_')
        filename = f'/tmp/{tmp}.docx'
        doc.save(filename)
        with open(filename, 'rb') as final_file:
            base64_string = base64.b64encode(final_file.read()).decode('utf-8')

        self.report_docx = base64_string
        self.report_docx_filename = title + '.docx'

        return {
            "type": "ir.actions.act_url",
            "url": f"/web/content/vit.product_value_analysis/{self.id}/report_docx/{self.report_docx_filename}",
            "target": "self",
        }

    def generate_output_html(self):
        self.output_html = self.md_to_html(
            self.json_to_markdown(
                json.loads(self.clean_md(self.output)), level=2, max_level=3
            )
        )

    def action_generate_market_mapping(self):
        self.market_mapper_ids.active=False

        mm = self.env['vit.market_mapper'].create({
            'name':'/',
            'gpt_model_id': self.gpt_model_id.id,
            'partner_id': self.partner_id.id,
            'lang_id': self.lang_id.id,
            'product_value_analysis_id': self.id,
        })
        mm._get_input()
        mm.action_generate()